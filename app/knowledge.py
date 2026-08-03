from __future__ import annotations

import json
import math
import re
import threading
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader

try:
    import psycopg
except ImportError:  # pragma: no cover
    psycopg = None


TOKEN_RE = re.compile(r"[A-Za-z0-9][A-Za-z0-9_'-]{1,}")
MATERIAL_TYPES = {"course", "approved_external"}
REPOSITORY_SCOPES = {"course", "admin_private", "shared"}


@dataclass
class Chunk:
    source: str
    chunk_index: int
    content: str
    class_id: str = ""
    material_type: str = "course"
    display_source: str = ""
    repository_scope: str = ""


@dataclass
class RetrievedChunk:
    source: str
    content: str
    score: float
    class_id: str = ""
    material_type: str = "course"


def _normalise_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        return data.decode("utf-8", errors="replace")
    if suffix == ".pdf":
        from io import BytesIO

        reader = PdfReader(BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        from io import BytesIO

        doc = Document(BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)
    raise ValueError("Unsupported file type. Use PDF, DOCX, TXT, MD or CSV.")


def make_chunks(
    text: str,
    source: str,
    target_words: int = 420,
    overlap_words: int = 60,
    *,
    class_id: str = "",
    material_type: str = "course",
    display_source: str = "",
    repository_scope: str = "",
) -> list[Chunk]:
    clean = _normalise_space(text)
    if not clean:
        return []
    material_type = material_type if material_type in MATERIAL_TYPES else "course"
    repository_scope = repository_scope if repository_scope in REPOSITORY_SCOPES else ("course" if class_id else "admin_private")
    words = clean.split()
    chunks: list[Chunk] = []
    start = 0
    index = 0
    while start < len(words):
        end = min(start + target_words, len(words))
        content = " ".join(words[start:end]).strip()
        if content:
            chunks.append(
                Chunk(
                    source=source,
                    chunk_index=index,
                    content=content,
                    class_id=class_id,
                    material_type=material_type,
                    display_source=display_source or source,
                    repository_scope=repository_scope,
                )
            )
            index += 1
        if end >= len(words):
            break
        start = max(end - overlap_words, start + 1)
    return chunks


class KnowledgeStore:
    def __init__(self, *, database_url: str, storage_dir: Path) -> None:
        self.database_url = database_url
        self.storage_file = storage_dir / "knowledge.json"
        self._lock = threading.RLock()
        self._memory: list[Chunk] = []
        self._use_postgres = bool(database_url and psycopg is not None)
        self.initialise()

    def _connect(self):
        if not self._use_postgres:
            raise RuntimeError("PostgreSQL is not configured")
        return psycopg.connect(self.database_url)

    def initialise(self) -> None:
        if self._use_postgres:
            with self._connect() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        """
                        CREATE TABLE IF NOT EXISTS ai_tutor_knowledge_chunks (
                            id BIGSERIAL PRIMARY KEY,
                            source TEXT NOT NULL,
                            chunk_index INTEGER NOT NULL,
                            content TEXT NOT NULL,
                            class_id TEXT NOT NULL DEFAULT '',
                            material_type TEXT NOT NULL DEFAULT 'course',
                            display_source TEXT NOT NULL DEFAULT '',
                            repository_scope TEXT NOT NULL DEFAULT '',
                            created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                            UNIQUE(source, chunk_index)
                        )
                        """
                    )
                    cur.execute("ALTER TABLE ai_tutor_knowledge_chunks ADD COLUMN IF NOT EXISTS class_id TEXT NOT NULL DEFAULT ''")
                    cur.execute("ALTER TABLE ai_tutor_knowledge_chunks ADD COLUMN IF NOT EXISTS material_type TEXT NOT NULL DEFAULT 'course'")
                    cur.execute("ALTER TABLE ai_tutor_knowledge_chunks ADD COLUMN IF NOT EXISTS display_source TEXT NOT NULL DEFAULT ''")
                    cur.execute("ALTER TABLE ai_tutor_knowledge_chunks ADD COLUMN IF NOT EXISTS repository_scope TEXT NOT NULL DEFAULT ''")
                    # Every historical class-less source is an administrator repository item.
                    # This migration prevents legacy uploads from leaking into lecturer courses.
                    cur.execute("""UPDATE ai_tutor_knowledge_chunks
                                   SET repository_scope = CASE WHEN class_id = '' THEN 'admin_private' ELSE 'course' END
                                   WHERE repository_scope IS NULL OR repository_scope = ''""")
                    cur.execute("CREATE INDEX IF NOT EXISTS idx_ai_tutor_knowledge_scope ON ai_tutor_knowledge_chunks(class_id, material_type, repository_scope)")
                conn.commit()
            return

        if self.storage_file.exists():
            try:
                raw = json.loads(self.storage_file.read_text(encoding="utf-8"))
                self._memory = [
                    Chunk(
                        source=str(item.get("source", "")),
                        chunk_index=int(item.get("chunk_index", 0)),
                        content=str(item.get("content", "")),
                        class_id=str(item.get("class_id", "")),
                        material_type=str(item.get("material_type", "course")),
                        display_source=str(item.get("display_source", item.get("source", ""))),
                        repository_scope=str(item.get("repository_scope") or ("course" if item.get("class_id") else "admin_private")),
                    )
                    for item in raw
                ]
            except (json.JSONDecodeError, TypeError, OSError, ValueError):
                self._memory = []

    def _save_local(self) -> None:
        payload = [chunk.__dict__ for chunk in self._memory]
        self.storage_file.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    def replace_source(self, source: str, chunks: Iterable[Chunk]) -> int:
        new_chunks = list(chunks)
        for chunk in new_chunks:
            if chunk.material_type not in MATERIAL_TYPES:
                chunk.material_type = "course"
            if chunk.repository_scope not in REPOSITORY_SCOPES:
                chunk.repository_scope = "course" if chunk.class_id else "admin_private"
        if self._use_postgres:
            with self._connect() as conn:
                with conn.cursor() as cur:
                    cur.execute("DELETE FROM ai_tutor_knowledge_chunks WHERE source = %s", (source,))
                    cur.executemany(
                        """INSERT INTO ai_tutor_knowledge_chunks
                           (source, chunk_index, content, class_id, material_type, display_source, repository_scope)
                           VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                        [
                            (
                                c.source,
                                c.chunk_index,
                                c.content,
                                c.class_id,
                                c.material_type if c.material_type in MATERIAL_TYPES else "course",
                                c.display_source or c.source,
                                c.repository_scope if c.repository_scope in REPOSITORY_SCOPES else ("course" if c.class_id else "admin_private"),
                            )
                            for c in new_chunks
                        ],
                    )
                conn.commit()
            return len(new_chunks)

        with self._lock:
            self._memory = [c for c in self._memory if c.source != source]
            self._memory.extend(new_chunks)
            self._save_local()
        return len(new_chunks)

    def all_chunks(self) -> list[Chunk]:
        if self._use_postgres:
            with self._connect() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        """SELECT source, chunk_index, content, class_id, material_type, display_source, repository_scope
                           FROM ai_tutor_knowledge_chunks ORDER BY source, chunk_index"""
                    )
                    rows = cur.fetchall()
            return [
                Chunk(
                    source=row[0],
                    chunk_index=row[1],
                    content=row[2],
                    class_id=row[3] or "",
                    material_type=row[4] or "course",
                    display_source=row[5] or row[0],
                    repository_scope=row[6] or ("course" if row[3] else "admin_private"),
                )
                for row in rows
            ]
        with self._lock:
            return list(self._memory)

    def list_sources(self, *, class_id: str | None = None, include_global: bool = False) -> list[dict[str, object]]:
        """List only sources belonging to the requested repository.

        Blank class ids are the private administrator repository. A lecturer or
        student course request can never include those rows, including legacy rows.
        """
        chunks = self.all_chunks()
        if class_id is None:
            pass
        elif class_id == "":
            chunks = [c for c in chunks if not c.class_id and c.repository_scope == "admin_private"]
        else:
            chunks = [c for c in chunks if c.class_id == class_id and c.repository_scope == "course"]
            if include_global:
                chunks.extend(c for c in self.all_chunks() if c.repository_scope == "shared" and not c.class_id)
        counts: Counter[tuple[str, str, str, str, str]] = Counter(
            (chunk.source, chunk.display_source or chunk.source, chunk.class_id, chunk.material_type, chunk.repository_scope)
            for chunk in chunks
        )
        return [
            {
                "source_id": source_id,
                "source": display_source,
                "chunks": count,
                "class_id": scope,
                "material_type": material_type,
                "repository_scope": repository_scope,
            }
            for (source_id, display_source, scope, material_type, repository_scope), count
            in sorted(counts.items(), key=lambda item: (item[0][1], item[0][2], item[0][3]))
        ]

    def source_metadata(self, source: str) -> dict[str, object] | None:
        source = str(source or "").strip()
        if not source:
            return None
        chunks = [chunk for chunk in self.all_chunks() if chunk.source == source]
        if not chunks:
            return None
        first = chunks[0]
        return {
            "source_id": source,
            "source": first.display_source or source,
            "class_id": first.class_id,
            "material_type": first.material_type,
            "repository_scope": first.repository_scope,
            "chunks": len(chunks),
        }

    def delete_source(self, source: str) -> int:
        """Delete an exact indexed source and return the number of removed chunks."""
        source = str(source or "").strip()
        if not source:
            return 0
        if self._use_postgres:
            with self._connect() as conn:
                with conn.cursor() as cur:
                    cur.execute("DELETE FROM ai_tutor_knowledge_chunks WHERE source = %s", (source,))
                    changed = max(int(cur.rowcount or 0), 0)
                conn.commit()
            return changed
        with self._lock:
            before = len(self._memory)
            self._memory = [chunk for chunk in self._memory if chunk.source != source]
            changed = before - len(self._memory)
            if changed:
                self._save_local()
            return changed

    def delete_admin_source(self, source: str) -> int:
        """Delete only a private administrator source, including legacy source names."""
        metadata = self.source_metadata(source)
        if not metadata:
            return 0
        if metadata.get("class_id") or metadata.get("repository_scope") != "admin_private":
            return 0
        return self.delete_source(source)

    def delete_course_document_sources(
        self, *, class_id: str, document_id: str, filename: str, document_type: str
    ) -> int:
        """Remove current and historical index aliases for one course document."""
        class_id = str(class_id or "").strip()
        document_id = str(document_id or "").strip()
        filename = str(filename or "").strip()
        document_type = str(document_type or "teaching_notes").strip()
        if not class_id or not document_id:
            return 0
        exact_sources = {
            f"{class_id}::{document_type}::{document_id}",
            f"{class_id}::{document_id}",
            f"{class_id}::{document_type}::{filename}",
            f"{class_id}::{filename}",
            filename,
        }
        if self._use_postgres:
            with self._connect() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        """DELETE FROM ai_tutor_knowledge_chunks
                           WHERE class_id = %s AND repository_scope = 'course'
                             AND (source = ANY(%s) OR source LIKE %s OR display_source = %s OR display_source LIKE %s)""",
                        (class_id, list(exact_sources), f"%::{document_id}", filename, f"{filename} •%"),
                    )
                    changed = max(int(cur.rowcount or 0), 0)
                conn.commit()
            return changed
        with self._lock:
            before = len(self._memory)
            self._memory = [
                chunk for chunk in self._memory
                if not (
                    chunk.class_id == class_id
                    and chunk.repository_scope == "course"
                    and (
                        chunk.source in exact_sources
                        or chunk.source.endswith(f"::{document_id}")
                        or (chunk.display_source or "") == filename
                        or (chunk.display_source or "").startswith(f"{filename} •")
                    )
                )
            ]
            changed = before - len(self._memory)
            if changed:
                self._save_local()
            return changed

    def retrieve(
        self,
        query: str,
        limit: int = 5,
        *,
        class_id: str = "",
        allowed_types: set[str] | None = None,
        include_global: bool = True,
    ) -> list[RetrievedChunk]:
        chunks = self.all_chunks()
        allowed_types = allowed_types or {"course", "approved_external"}
        chunks = [
            chunk
            for chunk in chunks
            if chunk.material_type in allowed_types
            and (
                (class_id and chunk.class_id == class_id and chunk.repository_scope == "course")
                or (include_global and chunk.repository_scope == "shared" and not chunk.class_id)
            )
        ]
        if not chunks or not query.strip():
            return []

        query_tokens = [t.lower() for t in TOKEN_RE.findall(query)]
        if not query_tokens:
            return []
        query_counts = Counter(query_tokens)

        doc_tokens: list[Counter[str]] = []
        document_frequency: Counter[str] = Counter()
        for chunk in chunks:
            counts = Counter(t.lower() for t in TOKEN_RE.findall(chunk.content))
            doc_tokens.append(counts)
            for token in counts:
                if token in query_counts:
                    document_frequency[token] += 1

        total_docs = len(chunks)
        scored: list[RetrievedChunk] = []
        query_lower = query.lower().strip()
        for chunk, counts in zip(chunks, doc_tokens):
            length_norm = max(sum(counts.values()), 1)
            score = 0.0
            for token, q_count in query_counts.items():
                tf = counts[token] / length_norm
                idf = math.log((total_docs + 1) / (document_frequency[token] + 1)) + 1.0
                score += tf * idf * (1 + math.log(q_count))
            if len(query_lower) >= 12 and query_lower in chunk.content.lower():
                score += 2.0
            if chunk.class_id == class_id and class_id:
                score *= 1.12
            if score > 0:
                scored.append(
                    RetrievedChunk(
                        source=chunk.display_source or chunk.source,
                        content=chunk.content,
                        score=score,
                        class_id=chunk.class_id,
                        material_type=chunk.material_type,
                    )
                )

        scored.sort(key=lambda item: item.score, reverse=True)
        return scored[:limit]
