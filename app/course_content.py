from __future__ import annotations

import json
import re
import sqlite3
import threading
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

try:
    import psycopg
    from psycopg.rows import dict_row
except ImportError:  # pragma: no cover
    psycopg = None
    dict_row = None


DOCUMENT_TYPES = {"teaching_notes", "course_outline", "recommended_reading"}
HEADING_PREFIX_RE = re.compile(
    r"^(chapter|unit|module|week|lesson|topic|section|part|theme)\s+[A-Za-z0-9IVXLC.-]+\b",
    re.IGNORECASE,
)
NUMBERED_HEADING_RE = re.compile(r"^(\d+(?:\.\d+){0,4})[.)]?\s+(.{2,140})$")
MARKDOWN_HEADING_RE = re.compile(r"^(#{1,5})\s+(.+)$")
BULLET_RE = re.compile(r"^\s*(?:[-•*]|\d+[.)]|[A-Za-z][.)])\s+")

WEEK_TITLE_RE = re.compile(r"^(?:week|wk)\s*([0-9]{1,2}|[IVXLC]+)\b(?:\s*[-:.)]\s*|\s+)?(.*)$", re.IGNORECASE)
TABLE_WEEK_HEADERS = {"week", "wk", "teaching week", "session", "period", "teaching period", "class period"}
TABLE_TOPIC_HEADERS = {"topic", "topics", "content", "unit", "theme", "lesson", "subject matter", "course content"}
TABLE_PREPARATION_HEADERS = {
    "student preparation", "student's preparation", "student’s preparation", "preparation",
    "student preparation/activity", "reading", "readings", "student task", "student tasks",
}
TABLE_ACTIVITY_HEADERS = {"activities", "learning activities", "weekly activities", "student activities", "learning activity"}
TABLE_SUBUNIT_HEADERS = {"subunit", "sub-unit", "subtopic", "sub-topic", "outline"}
NUMBER_WORDS = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6,
    "seven": 7, "eight": 8, "nine": 9, "ten": 10, "eleven": 11,
    "twelve": 12, "thirteen": 13, "fourteen": 14, "fifteen": 15,
    "sixteen": 16, "seventeen": 17, "eighteen": 18, "nineteen": 19, "twenty": 20,
}


def _cell_items(value: str) -> list[str]:
    raw = (value or "").replace("\r", "\n")
    parts: list[str] = []
    for line in raw.split("\n"):
        for item in re.split(r"\s*[;•]\s*", line):
            clean = BULLET_RE.sub("", _clean_line(item)).strip(" -•*\t")
            if clean and clean not in parts:
                parts.append(clean[:500])
    return parts


def _cell_paragraph_items(cell: Any) -> list[str]:
    """Preserve paragraph boundaries inside a Word table cell.

    Many course outlines store the main topic and every subtopic in separate
    paragraphs, even when the paragraph styles are inconsistent. Using
    ``cell.text`` followed by whitespace normalisation destroys those
    boundaries and collapses an entire semester into one line.
    """
    items: list[str] = []
    for paragraph in getattr(cell, "paragraphs", []):
        paragraph_text = _clean_line(paragraph.text)
        for raw_item in re.split(r"\s*[;•]\s*", paragraph_text):
            text = BULLET_RE.sub("", _clean_line(raw_item)).strip(" -•*\t")
            if text and text not in items:
                items.append(text[:500])
    return items or _cell_items(getattr(cell, "text", ""))


def _roman_to_int(value: str) -> int | None:
    value = value.upper().strip()
    if not value or not re.fullmatch(r"[IVXLC]+", value):
        return None
    values = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100}
    total = 0
    previous = 0
    for character in reversed(value):
        current = values[character]
        total += -current if current < previous else current
        previous = max(previous, current)
    return total if 0 < total <= 52 else None


def _normalise_week_number(value: str) -> str:
    clean = _clean_line(value).strip(" .:-()[]")
    clean = re.sub(r"^(?:week|wk|period|session)\s*", "", clean, flags=re.IGNORECASE).strip(" .:-()[]")
    first = clean.split()[0].lower() if clean else ""
    if first.isdigit():
        return str(int(first))
    if first in NUMBER_WORDS:
        return str(NUMBER_WORDS[first])
    roman = _roman_to_int(first)
    return str(roman) if roman is not None else clean[:30]


def _course_title_from_tables(document: Document) -> str:
    for table in document.tables:
        for row in table.rows:
            values = [_clean_line(cell.text) for cell in row.cells]
            if len(values) < 2:
                continue
            label = values[0].lower().replace("/", " ")
            if "course code" in label or "course title" in label:
                return values[1][:200]
    return ""


def _table_week_sections(document: Document) -> list[tuple[int, str, list[str]]]:
    sections: list[tuple[int, str, list[str]]] = []
    for table_number, table in enumerate(document.tables, start=1):
        if not table.rows:
            continue
        header_cells = table.rows[0].cells
        headers = [_clean_line(cell.text).lower().strip() for cell in header_cells]

        def column(names: set[str]) -> int | None:
            for idx, header in enumerate(headers):
                if header in names or any(name in header for name in names):
                    return idx
            return None

        week_idx = column(TABLE_WEEK_HEADERS)
        topic_idx = column(TABLE_TOPIC_HEADERS)
        preparation_idx = column(TABLE_PREPARATION_HEADERS)
        activity_idx = column(TABLE_ACTIVITY_HEADERS)
        explicit_subunit_idx = column(TABLE_SUBUNIT_HEADERS)

        # A week/period table must contain both a period marker and a topic/content column.
        if week_idx is not None and topic_idx is not None:
            current_week = ""
            for row in table.rows[1:]:
                cells = row.cells
                week_items = _cell_paragraph_items(cells[week_idx]) if week_idx < len(cells) else []
                if week_items:
                    current_week = week_items[0]
                if not current_week:
                    continue

                topic_items = _cell_paragraph_items(cells[topic_idx]) if topic_idx < len(cells) else []
                if not topic_items:
                    continue
                main_topic = topic_items[0]
                week_number = _normalise_week_number(current_week)
                week_label = f"Week {week_number}" if week_number else f"Week {current_week}"
                title = f"{week_label}: {main_topic}" if main_topic else week_label

                preparation_items = (
                    _cell_paragraph_items(cells[preparation_idx])
                    if preparation_idx is not None and preparation_idx < len(cells)
                    else []
                )
                activity_items = (
                    _cell_paragraph_items(cells[activity_idx])
                    if activity_idx is not None and activity_idx < len(cells)
                    else []
                )
                content_lines = [f"Main topic: {main_topic}"]
                if preparation_items or activity_items:
                    content_lines.append("Student preparation and activities:")
                    content_lines.extend(f"- {item}" for item in [*preparation_items, *activity_items])
                sections.append((1, title[:200], content_lines))

                subunits = topic_items[1:]
                # In many outlines an “Activities” column actually lists the lesson's
                # subtopics. Use those items as subunits when the topic cell contains
                # only a main heading, while still showing them as weekly activities.
                if len(topic_items) == 1:
                    for item in activity_items:
                        if item not in subunits:
                            subunits.append(item)
                if explicit_subunit_idx is not None and explicit_subunit_idx != topic_idx and explicit_subunit_idx < len(cells):
                    for item in _cell_paragraph_items(cells[explicit_subunit_idx]):
                        if item not in subunits:
                            subunits.append(item)
                for subunit in subunits:
                    if subunit.lower() == main_topic.lower():
                        continue
                    sections.append((2, subunit[:200], [f"Part of {title}.", subunit]))
            continue

        # Keep useful non-week tables with meaningful labels, but skip decorative logo/header tables.
        rows = [[_clean_line(cell.text) for cell in row.cells] for row in table.rows]
        rows = [row for row in rows if any(row)]
        if not rows:
            continue
        first_column = {row[0].lower() for row in rows if row and row[0]}
        if any("course code" in value or "lecturer" in value for value in first_column):
            section_title = "Course information"
        elif any(value in {"attendance", "assignments", "academic dishonesty", "class participation"} for value in first_column):
            section_title = "Course policies"
        else:
            # Single-cell institutional headers and layout-only tables are not learning sections.
            meaningful_cells = [value for row in rows for value in row if value]
            if len(rows) <= 1 or len(meaningful_cells) <= 1:
                continue
            section_title = f"Supporting table {table_number}"
        table_lines = [" | ".join(value for value in row if value) for row in rows]
        sections.append((1, section_title, table_lines))
    return sections

def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _iso(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.astimezone(timezone.utc).isoformat()
    return str(value)


def _json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))


def _safe_json(value: Any, default: Any) -> Any:
    if isinstance(value, (dict, list)):
        return value
    if value is None:
        return default
    try:
        return json.loads(value)
    except (TypeError, ValueError, json.JSONDecodeError):
        return default


def _clean_line(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _looks_like_heading(line: str) -> tuple[int, str] | None:
    clean = _clean_line(line)
    if not clean or len(clean) > 160:
        return None
    md = MARKDOWN_HEADING_RE.match(clean)
    if md:
        return min(len(md.group(1)), 5), _clean_line(md.group(2))
    numbered = NUMBERED_HEADING_RE.match(clean)
    if numbered:
        number_parts = numbered.group(1).split(".")
        while len(number_parts) > 1 and number_parts[-1] == "0":
            number_parts.pop()
        level = min(max(1, len(number_parts)), 5)
        return level, _clean_line(f"{numbered.group(1)} {numbered.group(2)}")
    if HEADING_PREFIX_RE.match(clean):
        return 1 if clean.lower().startswith(("chapter", "unit", "module", "part")) else 2, clean
    words = clean.split()
    if 2 <= len(words) <= 12 and clean.isupper() and any(ch.isalpha() for ch in clean):
        return 1, clean.title()
    if clean.endswith(":") and 1 <= len(words) <= 10:
        return 2, clean.rstrip(":")
    return None


def _extract_docx_blocks(data: bytes) -> tuple[str, list[tuple[int, str, list[str]]]]:
    document = Document(BytesIO(data))
    title = _course_title_from_tables(document)
    sections: list[tuple[int, str, list[str]]] = []
    current_level = 1
    current_title = "Complete document"
    current_lines: list[str] = []

    def flush(include_empty: bool = False) -> None:
        nonlocal current_lines
        content = "\n".join(line for line in current_lines if line.strip()).strip()
        if content or (include_empty and current_title != "Complete document"):
            sections.append((current_level, current_title, current_lines[:]))
        current_lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style_name == "title" and not title:
            title = text
            continue
        heading_match = re.search(r"heading\s*([1-5])", style_name)
        detected = _looks_like_heading(text)
        list_context = any(term in current_title.lower() for term in ("objective", "outcome", "recommended reading", "bibliograph", "reference"))
        # Introductory sentences such as “This course enables students to:” belong
        # inside the objective section. A new numbered heading, such as “3.0
        # Course Outcomes”, must still close the previous section.
        numbered_in_context = NUMBERED_HEADING_RE.match(text)
        if detected and list_context and not heading_match:
            # A simple “1. Explain …” line is normally a list item inside
            # objectives/readings. A decimal heading such as “3.0 Course
            # Outcomes” is a genuine new section and must remain a heading.
            if not numbered_in_context or "." not in numbered_in_context.group(1):
                detected = None
        if heading_match or detected:
            flush(include_empty=True)
            if detected and NUMBERED_HEADING_RE.match(text):
                current_level, current_title = detected
            elif heading_match:
                current_level = int(heading_match.group(1))
                current_title = text
            else:
                current_level, current_title = detected or (2, text)
        else:
            current_lines.append(text)
    flush(include_empty=True)
    sections.extend(_table_week_sections(document))
    return title, sections


def _extract_text_blocks(filename: str, data: bytes) -> tuple[str, list[tuple[int, str, list[str]]]]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return _extract_docx_blocks(data)
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(data))
        raw = "\n\n".join((page.extract_text() or "") for page in reader.pages)
    elif suffix in {".txt", ".md", ".csv"}:
        raw = data.decode("utf-8", errors="replace")
    else:
        raise ValueError("Unsupported file type. Use PDF, DOCX, TXT, MD or CSV.")

    lines = [line.rstrip() for line in raw.replace("\r", "\n").split("\n")]
    title = ""
    sections: list[tuple[int, str, list[str]]] = []
    current_level = 1
    current_title = "Complete document"
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        content = "\n".join(line for line in current_lines if line.strip()).strip()
        if content:
            sections.append((current_level, current_title, current_lines[:]))
        current_lines = []

    for raw_line in lines:
        line = _clean_line(raw_line)
        if not line:
            if current_lines and current_lines[-1] != "":
                current_lines.append("")
            continue
        heading = _looks_like_heading(line)
        if heading:
            if not title and not sections and not current_lines and heading[0] == 1:
                title = heading[1]
                current_title = title
                continue
            flush()
            current_level, current_title = heading
        else:
            current_lines.append(line)
    flush()
    return title, sections


def _extract_list(lines: list[str]) -> list[str]:
    items: list[str] = []
    for line in lines:
        clean = BULLET_RE.sub("", _clean_line(line)).strip(" -•*\t")
        if not clean or len(clean) < 4:
            continue
        if clean.lower().startswith(("students should", "by the end", "upon completion", "this course enables")) and ":" in clean:
            clean = clean.split(":", 1)[1].strip()
        if not clean:
            continue
        if clean and clean not in items:
            items.append(clean[:500])
    return items[:40]


@dataclass
class ParsedSection:
    title: str
    level: int
    content: str
    position: int
    parent_position: int | None
    section_path: str


@dataclass
class ParsedDocument:
    title: str
    sections: list[ParsedSection]
    objectives: list[str]
    recommended_readings: list[str]
    weekly_topics: list[str]


def parse_document(filename: str, data: bytes) -> ParsedDocument:
    detected_title, raw_sections = _extract_text_blocks(filename, data)
    title = detected_title or Path(filename).stem.replace("_", " ").replace("-", " ").strip().title()
    stack: list[tuple[int, int, str]] = []
    parsed: list[ParsedSection] = []
    objectives: list[str] = []
    readings: list[str] = []
    weekly_topics: list[str] = []

    for position, (level, section_title, section_lines) in enumerate(raw_sections):
        section_title = _clean_line(section_title) or f"Section {position + 1}"
        level = max(1, min(int(level or 1), 5))
        while stack and stack[-1][0] >= level:
            stack.pop()
        parent_position = stack[-1][1] if stack else None
        path_titles = [item[2] for item in stack] + [section_title]
        content = "\n".join(section_lines).strip()
        parsed.append(
            ParsedSection(
                title=section_title,
                level=level,
                content=content,
                position=position,
                parent_position=parent_position,
                section_path=" > ".join(path_titles),
            )
        )
        stack.append((level, position, section_title))
        lower_title = section_title.lower()
        if WEEK_TITLE_RE.match(section_title):
            weekly_topics.append(section_title[:300])
        if any(term in lower_title for term in ("objective", "outcome")):
            objectives.extend(_extract_list(section_lines))
        if any(term in lower_title for term in ("recommended reading", "reading list", "references", "bibliography")):
            readings.extend(_extract_list(section_lines))

    if not parsed:
        parsed = [ParsedSection("Complete document", 1, "", 0, None, "Complete document")]
    objectives = list(dict.fromkeys(objectives))[:30]
    readings = list(dict.fromkeys(readings))[:60]
    weekly_topics = list(dict.fromkeys(weekly_topics))[:40]
    return ParsedDocument(title=title[:200], sections=parsed, objectives=objectives, recommended_readings=readings, weekly_topics=weekly_topics)


class CourseContentStore:
    """Persistent document hierarchy for lecturer-controlled course portals."""

    def __init__(self, *, database_url: str, storage_dir: Path) -> None:
        self.database_url = database_url
        self._use_postgres = bool(database_url and psycopg is not None)
        self.sqlite_path = storage_dir / "ai_tutor_course_content.sqlite3"
        self._lock = threading.RLock()
        self.initialise()

    def _pg(self):
        if not self._use_postgres:
            raise RuntimeError("PostgreSQL is not configured")
        return psycopg.connect(self.database_url, row_factory=dict_row)

    def _sqlite(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.sqlite_path, timeout=20, check_same_thread=False)
        conn.row_factory = sqlite3.Row
        return conn

    def initialise(self) -> None:
        if self._use_postgres:
            statements = [
                """
                CREATE TABLE IF NOT EXISTS ai_tutor_course_documents (
                    id TEXT PRIMARY KEY,
                    class_id TEXT NOT NULL,
                    uploader_id TEXT NOT NULL,
                    title TEXT NOT NULL,
                    filename TEXT NOT NULL,
                    document_type TEXT NOT NULL,
                    objectives JSONB NOT NULL DEFAULT '[]'::jsonb,
                    recommended_readings JSONB NOT NULL DEFAULT '[]'::jsonb,
                    weekly_topics JSONB NOT NULL DEFAULT '[]'::jsonb,
                    created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
                )
                """,
                """
                CREATE TABLE IF NOT EXISTS ai_tutor_course_sections (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL REFERENCES ai_tutor_course_documents(id) ON DELETE CASCADE,
                    class_id TEXT NOT NULL,
                    parent_id TEXT,
                    title TEXT NOT NULL,
                    section_path TEXT NOT NULL DEFAULT '',
                    level INTEGER NOT NULL DEFAULT 1,
                    position INTEGER NOT NULL DEFAULT 0,
                    content TEXT NOT NULL DEFAULT '',
                    created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
                )
                """,
                "ALTER TABLE ai_tutor_course_documents ADD COLUMN IF NOT EXISTS weekly_topics JSONB NOT NULL DEFAULT '[]'::jsonb",
                "CREATE INDEX IF NOT EXISTS idx_ai_tutor_course_docs_class ON ai_tutor_course_documents(class_id, created_at)",
                "CREATE INDEX IF NOT EXISTS idx_ai_tutor_course_sections_class ON ai_tutor_course_sections(class_id, document_id, position)",
            ]
            with self._pg() as conn, conn.cursor() as cur:
                for statement in statements:
                    cur.execute(statement)
                conn.commit()
            return
        schema = """
        CREATE TABLE IF NOT EXISTS ai_tutor_course_documents (
            id TEXT PRIMARY KEY, class_id TEXT NOT NULL, uploader_id TEXT NOT NULL,
            title TEXT NOT NULL, filename TEXT NOT NULL, document_type TEXT NOT NULL,
            objectives TEXT NOT NULL DEFAULT '[]', recommended_readings TEXT NOT NULL DEFAULT '[]', weekly_topics TEXT NOT NULL DEFAULT '[]',
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS ai_tutor_course_sections (
            id TEXT PRIMARY KEY, document_id TEXT NOT NULL, class_id TEXT NOT NULL,
            parent_id TEXT, title TEXT NOT NULL, section_path TEXT NOT NULL DEFAULT '',
            level INTEGER NOT NULL DEFAULT 1, position INTEGER NOT NULL DEFAULT 0,
            content TEXT NOT NULL DEFAULT '', created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_ai_tutor_course_docs_class ON ai_tutor_course_documents(class_id, created_at);
        CREATE INDEX IF NOT EXISTS idx_ai_tutor_course_sections_class ON ai_tutor_course_sections(class_id, document_id, position);
        """
        with self._lock, self._sqlite() as conn:
            conn.executescript(schema)
            columns = {row[1] for row in conn.execute("PRAGMA table_info(ai_tutor_course_documents)").fetchall()}
            if "weekly_topics" not in columns:
                conn.execute("ALTER TABLE ai_tutor_course_documents ADD COLUMN weekly_topics TEXT NOT NULL DEFAULT '[]'")
            conn.commit()

    def ingest_document(
        self,
        *,
        class_id: str,
        uploader_id: str,
        filename: str,
        document_type: str,
        data: bytes,
    ) -> dict[str, Any]:
        if document_type not in DOCUMENT_TYPES:
            document_type = "teaching_notes"
        parsed = parse_document(filename, data)
        document_id = str(uuid.uuid4())
        now = _utcnow()
        section_ids = {section.position: str(uuid.uuid4()) for section in parsed.sections}

        if self._use_postgres:
            with self._pg() as conn, conn.cursor() as cur:
                cur.execute(
                    "DELETE FROM ai_tutor_course_documents WHERE class_id=%s AND filename=%s AND document_type=%s",
                    (class_id, filename, document_type),
                )
                cur.execute(
                    """INSERT INTO ai_tutor_course_documents
                       (id,class_id,uploader_id,title,filename,document_type,objectives,recommended_readings,weekly_topics,created_at)
                       VALUES(%s,%s,%s,%s,%s,%s,%s::jsonb,%s::jsonb,%s::jsonb,%s)""",
                    (
                        document_id, class_id, uploader_id, parsed.title, filename, document_type,
                        _json(parsed.objectives), _json(parsed.recommended_readings), _json(parsed.weekly_topics), now,
                    ),
                )
                for section in parsed.sections:
                    cur.execute(
                        """INSERT INTO ai_tutor_course_sections
                           (id,document_id,class_id,parent_id,title,section_path,level,position,content,created_at)
                           VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (
                            section_ids[section.position], document_id, class_id,
                            section_ids.get(section.parent_position) if section.parent_position is not None else None,
                            section.title, section.section_path, section.level, section.position,
                            section.content, now,
                        ),
                    )
                conn.commit()
        else:
            with self._lock, self._sqlite() as conn:
                old_ids = [row[0] for row in conn.execute(
                    "SELECT id FROM ai_tutor_course_documents WHERE class_id=? AND filename=? AND document_type=?",
                    (class_id, filename, document_type),
                ).fetchall()]
                for old_id in old_ids:
                    conn.execute("DELETE FROM ai_tutor_course_sections WHERE document_id=?", (old_id,))
                    conn.execute("DELETE FROM ai_tutor_course_documents WHERE id=?", (old_id,))
                conn.execute(
                    """INSERT INTO ai_tutor_course_documents
                       (id,class_id,uploader_id,title,filename,document_type,objectives,recommended_readings,weekly_topics,created_at)
                       VALUES(?,?,?,?,?,?,?,?,?,?)""",
                    (
                        document_id, class_id, uploader_id, parsed.title, filename, document_type,
                        _json(parsed.objectives), _json(parsed.recommended_readings), _json(parsed.weekly_topics), now.isoformat(),
                    ),
                )
                for section in parsed.sections:
                    conn.execute(
                        """INSERT INTO ai_tutor_course_sections
                           (id,document_id,class_id,parent_id,title,section_path,level,position,content,created_at)
                           VALUES(?,?,?,?,?,?,?,?,?,?)""",
                        (
                            section_ids[section.position], document_id, class_id,
                            section_ids.get(section.parent_position) if section.parent_position is not None else None,
                            section.title, section.section_path, section.level, section.position,
                            section.content, now.isoformat(),
                        ),
                    )
                conn.commit()
        result = self.get_document(document_id)
        return result or {}

    @staticmethod
    def _document_public(row: dict[str, Any]) -> dict[str, Any]:
        return {
            "id": str(row.get("id", "")),
            "class_id": str(row.get("class_id", "")),
            "title": str(row.get("title", "")),
            "filename": str(row.get("filename", "")),
            "document_type": str(row.get("document_type", "teaching_notes")),
            "objectives": _safe_json(row.get("objectives"), []),
            "recommended_readings": _safe_json(row.get("recommended_readings"), []),
            "weekly_topics": _safe_json(row.get("weekly_topics"), []),
            "created_at": _iso(row.get("created_at")),
        }

    @staticmethod
    def _section_public(row: dict[str, Any], include_content: bool = False) -> dict[str, Any]:
        result = {
            "id": str(row.get("id", "")),
            "document_id": str(row.get("document_id", "")),
            "class_id": str(row.get("class_id", "")),
            "parent_id": str(row.get("parent_id", "") or ""),
            "title": str(row.get("title", "")),
            "section_path": str(row.get("section_path", "")),
            "level": int(row.get("level", 1) or 1),
            "position": int(row.get("position", 0) or 0),
            "created_at": _iso(row.get("created_at")),
        }
        if include_content:
            result["content"] = str(row.get("content", ""))
        return result

    def get_document(self, document_id: str) -> dict[str, Any] | None:
        if self._use_postgres:
            with self._pg() as conn, conn.cursor() as cur:
                cur.execute("SELECT * FROM ai_tutor_course_documents WHERE id=%s", (document_id,))
                row = cur.fetchone()
                if not row:
                    return None
                result = self._document_public(dict(row))
                cur.execute("SELECT * FROM ai_tutor_course_sections WHERE document_id=%s ORDER BY position", (document_id,))
                result["sections"] = [self._section_public(dict(item)) for item in cur.fetchall()]
                return result
        with self._lock, self._sqlite() as conn:
            row = conn.execute("SELECT * FROM ai_tutor_course_documents WHERE id=?", (document_id,)).fetchone()
            if not row:
                return None
            result = self._document_public(dict(row))
            sections = conn.execute("SELECT * FROM ai_tutor_course_sections WHERE document_id=? ORDER BY position", (document_id,)).fetchall()
            result["sections"] = [self._section_public(dict(item)) for item in sections]
            return result

    def list_structure(self, class_id: str) -> list[dict[str, Any]]:
        if self._use_postgres:
            with self._pg() as conn, conn.cursor() as cur:
                cur.execute("SELECT * FROM ai_tutor_course_documents WHERE class_id=%s ORDER BY created_at, title", (class_id,))
                documents = [self._document_public(dict(row)) for row in cur.fetchall()]
                cur.execute("SELECT * FROM ai_tutor_course_sections WHERE class_id=%s ORDER BY document_id, position", (class_id,))
                sections = [self._section_public(dict(row)) for row in cur.fetchall()]
        else:
            with self._lock, self._sqlite() as conn:
                documents = [self._document_public(dict(row)) for row in conn.execute(
                    "SELECT * FROM ai_tutor_course_documents WHERE class_id=? ORDER BY created_at, title", (class_id,)
                ).fetchall()]
                sections = [self._section_public(dict(row)) for row in conn.execute(
                    "SELECT * FROM ai_tutor_course_sections WHERE class_id=? ORDER BY document_id, position", (class_id,)
                ).fetchall()]
        by_document: dict[str, list[dict[str, Any]]] = {}
        for section in sections:
            by_document.setdefault(section["document_id"], []).append(section)
        for document in documents:
            document["sections"] = by_document.get(document["id"], [])
        return documents

    def get_section(self, section_id: str) -> dict[str, Any] | None:
        sql = """
        SELECT s.*, d.title AS document_title, d.filename, d.document_type,
               d.objectives, d.recommended_readings, d.weekly_topics
        FROM ai_tutor_course_sections s
        JOIN ai_tutor_course_documents d ON d.id=s.document_id
        WHERE s.id={placeholder}
        """
        if self._use_postgres:
            with self._pg() as conn, conn.cursor() as cur:
                cur.execute(sql.format(placeholder="%s"), (section_id,))
                row = cur.fetchone()
                if not row:
                    return None
                data = dict(row)
        else:
            with self._lock, self._sqlite() as conn:
                row = conn.execute(sql.format(placeholder="?"), (section_id,)).fetchone()
                if not row:
                    return None
                data = dict(row)
        result = self._section_public(data, include_content=True)
        result.update({
            "document_title": str(data.get("document_title", "")),
            "filename": str(data.get("filename", "")),
            "document_type": str(data.get("document_type", "teaching_notes")),
            "objectives": _safe_json(data.get("objectives"), []),
            "recommended_readings": _safe_json(data.get("recommended_readings"), []),
            "weekly_topics": _safe_json(data.get("weekly_topics"), []),
        })
        return result

    def weekly_plan(self, class_id: str, fallback_topics: list[str] | None = None, fallback_outcomes: list[str] | None = None) -> list[dict[str, Any]]:
        documents = self.list_structure(class_id)
        entries: list[dict[str, Any]] = []
        seen: set[str] = set()
        for document in documents:
            if document.get("document_type") != "course_outline":
                continue
            sections = document.get("sections", [])
            for section in sections:
                title = str(section.get("title", "")).strip()
                if not WEEK_TITLE_RE.match(title):
                    continue
                key = title.lower()
                if key in seen:
                    continue
                seen.add(key)
                children = [
                    {
                        "id": str(child.get("id", "")),
                        "title": str(child.get("title", "")),
                        "section_path": str(child.get("section_path", child.get("title", ""))),
                    }
                    for child in sections
                    if str(child.get("parent_id", "")) == str(section.get("id", "")) and str(child.get("title", "")).strip()
                ]
                full_section = self.get_section(str(section.get("id", ""))) or {}
                preparation: list[str] = []
                collecting_preparation = False
                for raw_line in str(full_section.get("content", "")).splitlines():
                    line = _clean_line(raw_line)
                    if not line:
                        continue
                    if line.lower().startswith("student preparation and activities"):
                        collecting_preparation = True
                        continue
                    if collecting_preparation:
                        item = BULLET_RE.sub("", line).strip(" -•*\t")
                        if item and item not in preparation:
                            preparation.append(item[:500])
                entries.append({
                    "id": str(section.get("id", "")), "title": title,
                    "section_path": str(section.get("section_path", title)),
                    "subunits": children, "preparation": preparation[:20], "generated": False,
                })
        if entries:
            return entries[:40]
        topics = [str(item).strip() for item in (fallback_topics or []) if str(item).strip()]
        if not topics:
            topics = [str(item).strip() for item in (fallback_outcomes or []) if str(item).strip()]
        return [
            {
                "id": f"virtual:{class_id}:{index}",
                "title": topic if WEEK_TITLE_RE.match(topic) else f"Week {index + 1}: {topic}",
                "section_path": topic, "subunits": [], "preparation": [], "generated": True,
            }
            for index, topic in enumerate(topics[:40])
        ]

    def recommended_context(self, class_id: str, limit_chars: int = 14000) -> tuple[str, list[str]]:
        documents = self.list_structure(class_id)
        reading_doc_ids = {doc["id"] for doc in documents if doc["document_type"] == "recommended_reading"}
        if not reading_doc_ids:
            readings = []
            for doc in documents:
                readings.extend(doc.get("recommended_readings", []))
            return "\n".join(f"- {item}" for item in readings)[:limit_chars], []
        parts: list[str] = []
        sources: list[str] = []
        for doc in documents:
            if doc["id"] not in reading_doc_ids:
                continue
            sources.append(doc["filename"])
            for section in doc.get("sections", []):
                full = self.get_section(section["id"])
                if full and full.get("content"):
                    parts.append(f"[{doc['filename']} • {section['section_path']}]\n{full['content']}")
                    if sum(len(p) for p in parts) >= limit_chars:
                        return "\n\n".join(parts)[:limit_chars], sources
        return "\n\n".join(parts)[:limit_chars], sources

    def delete_document(self, document_id: str, class_id: str) -> bool:
        if self._use_postgres:
            with self._pg() as conn, conn.cursor() as cur:
                cur.execute("DELETE FROM ai_tutor_course_documents WHERE id=%s AND class_id=%s", (document_id, class_id))
                changed = cur.rowcount > 0
                conn.commit()
                return changed
        with self._lock, self._sqlite() as conn:
            row = conn.execute("SELECT 1 FROM ai_tutor_course_documents WHERE id=? AND class_id=?", (document_id, class_id)).fetchone()
            if not row:
                return False
            conn.execute("DELETE FROM ai_tutor_course_sections WHERE document_id=?", (document_id,))
            conn.execute("DELETE FROM ai_tutor_course_documents WHERE id=?", (document_id,))
            conn.commit()
            return True
