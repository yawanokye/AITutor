import base64
import os
import tempfile
import uuid
from io import BytesIO

os.environ["DEMO_MODE"] = "true"
os.environ["ADMIN_KEY"] = "test-admin-bootstrap"
os.environ["AUTH_SECRET"] = "test-auth-secret-long-enough"
os.environ["DATABASE_URL"] = ""
os.environ["VISUAL_PLAN_ENABLED"] = "true"
os.environ["REQUIRE_LOGIN_FOR_AI"] = "false"
os.environ["RATE_LIMIT_PER_MINUTE"] = "500"
os.environ["STORAGE_DIR"] = tempfile.mkdtemp(prefix="ai-tutor-v5-tests-")
os.environ["ALLOW_PUBLIC_TEACHER_REGISTRATION"] = "false"

from docx import Document
from fastapi.testclient import TestClient

from app.main import (
    _audio_upload_extension,
    _base_media_type,
    _extract_response_text,
    _normalise_visual_plan,
    _normalise_practice_evaluation,
    app,
    knowledge,
)
from app.schemas import PracticeEvaluation, SpeechRequest, VisualPlan
from app.knowledge import make_chunks


client = TestClient(app)
ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)
_ADMIN = None


def headers(auth):
    return {"Authorization": f"Bearer {auth['access_token']}"}


def unique_email(prefix):
    return f"{prefix}-{uuid.uuid4().hex[:12]}@example.com"


def admin_account():
    global _ADMIN
    if _ADMIN:
        return _ADMIN
    response = client.post(
        "/api/admin/bootstrap",
        json={
            "admin_key": "test-admin-bootstrap",
            "display_name": "Test Administrator",
            "email": unique_email("admin"),
            "password": "AdminStrong123!",
        },
    )
    assert response.status_code == 200, response.text
    _ADMIN = response.json()
    return _ADMIN


def lecturer_account(name="Test Lecturer"):
    admin = admin_account()
    email = unique_email("lecturer")
    temporary = "LecturerTemp123!"
    created = client.post(
        "/api/admin/lecturers",
        headers=headers(admin),
        json={"display_name": name, "email": email, "temporary_password": temporary},
    )
    assert created.status_code == 200, created.text
    login = client.post("/api/auth/login", json={"email": email, "password": temporary})
    assert login.status_code == 200, login.text
    return login.json(), created.json()


def student_account(name="Test Student"):
    response = client.post(
        "/api/auth/register",
        json={
            "display_name": name,
            "email": unique_email("student"),
            "password": "StudentStrong123!",
            "role": "student",
            "teacher_invite_code": "",
        },
    )
    assert response.status_code == 200, response.text
    return response.json()


def create_course(lecturer, *, name="Statistics 101", required=False, response_mode="student_choice", weekly_topics=None, outcomes=None):
    response = client.post(
        "/api/classes",
        headers=headers(lecturer),
        json={
            "name": name,
            "subject": "STA 101",
            "knowledge_mode": "course_only",
            "learning_outcomes": outcomes or [],
            "weekly_topics": weekly_topics or [],
            "recommended_readings": [],
            "tutor_instructions": "Use the uploaded lecturer notes as the main authority.",
            "practice_whiteboard_required": required,
            "practice_response_mode": "whiteboard" if required else response_mode,
        },
    )
    assert response.status_code == 200, response.text
    return response.json()


def outline_docx_bytes():
    output = BytesIO()
    document = Document()
    document.add_heading("STA 101 Detailed Course Outline", 0)
    document.add_heading("Course Objectives", 1)
    document.add_paragraph("1. Explain descriptive statistics")
    document.add_paragraph("2. Distinguish qualitative and quantitative data")
    document.add_heading("Recommended Reading", 1)
    document.add_paragraph("1. Author, A. Introductory Statistics")
    document.add_heading("1. Introduction to Statistics", 1)
    document.add_paragraph("Statistics concerns the collection, organisation, analysis and interpretation of data.")
    document.add_heading("1.1 Types of Data", 2)
    document.add_paragraph("Qualitative data describe categories. Quantitative data are numerical.")
    document.save(output)
    return output.getvalue()


def upload_outline(lecturer, classroom):
    response = client.post(
        "/api/materials/upload",
        headers=headers(lecturer),
        data={"class_id": classroom["id"], "document_type": "course_outline"},
        files=[(
            "files",
            ("statistics-outline.docx", outline_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        )],
    )
    assert response.status_code == 200, response.text
    return response.json()


def enrol(student, classroom):
    response = client.post(
        "/api/classes/join",
        headers=headers(student),
        json={"join_code": classroom["join_code"]},
    )
    assert response.status_code == 200, response.text
    return response.json()


def test_health_reports_v5_portal_build():
    response = client.get("/health")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "ok"
    assert data["version"] == "5.4.0"
    assert data["live_video_enabled"] is False
    assert data["institutional_mode"] is True
    assert data["course_lock_enabled"] is True
    assert data["visual_plan_enabled"] is True


def test_config_exposes_interactive_features():
    data = client.get("/api/config").json()
    assert data["demo_mode"] is True
    assert data["visual_plan_enabled"] is True
    assert data["interactive_practice_enabled"] is True
    assert data["work_check_enabled"] is True
    assert data["institutional_mode"] is True
    assert data["live_video_enabled"] is False


def test_index_contains_v5_role_portals_and_two_whiteboards():
    html = client.get("/").text
    for identifier in (
        'id="drawingCanvas"', 'id="practiceDrawingCanvas"', 'id="courseNavigatorPanel"',
        'id="courseStructureList"', 'id="documentType"', 'id="showAdminSetup"',
        'id="openDashboard"', 'id="classSelect"', 'id="outcomeSelect"', 'id="weekSelect"',
    ):
        assert identifier in html
    assert '/static/portal.js?v=5.4.0' in html
    assert '/static/practice_board.js?v=5.4.0' in html
    assert 'Administrator sign in' in html
    assert 'Lecturer sign in' in html
    assert 'Student sign in' in html
    assert 'Teacher invitation code' not in html
    assert 'openLiveVideo' not in html


def test_first_administrator_bootstrap_and_admin_dashboard():
    admin = admin_account()
    me = client.get("/api/auth/me", headers=headers(admin))
    assert me.status_code == 200
    assert me.json()["role"] == "admin"
    dashboard = client.get("/api/dashboard", headers=headers(admin))
    assert dashboard.status_code == 200
    assert dashboard.json()["role"] == "admin"
    assert "lecturers" in dashboard.json()


def test_second_administrator_bootstrap_is_blocked():
    admin_account()
    response = client.post(
        "/api/admin/bootstrap",
        json={
            "admin_key": "test-admin-bootstrap",
            "display_name": "Another Administrator",
            "email": unique_email("admin2"),
            "password": "AdminStrong123!",
        },
    )
    assert response.status_code == 409


def test_public_lecturer_registration_is_blocked():
    response = client.post(
        "/api/auth/register",
        json={
            "display_name": "Public Lecturer",
            "email": unique_email("public-lecturer"),
            "password": "StrongPass123!",
            "role": "teacher",
            "teacher_invite_code": "any-code",
        },
    )
    assert response.status_code == 403
    assert "administrator" in response.json()["detail"].lower()


def test_administrator_creates_lecturer_with_temporary_password():
    lecturer, created = lecturer_account("Dr Created Lecturer")
    assert lecturer["user"]["role"] == "teacher"
    assert created["temporary_password"]
    assert created["user"]["must_change_password"] is True
    listing = client.get("/api/admin/lecturers", headers=headers(admin_account()))
    assert listing.status_code == 200
    assert any(item["id"] == created["user"]["id"] for item in listing.json())


def test_lecturer_changes_temporary_password():
    lecturer, _ = lecturer_account()
    response = client.post(
        "/api/auth/change-password",
        headers=headers(lecturer),
        json={"current_password": "LecturerTemp123!", "new_password": "NewLecturerStrong123!"},
    )
    assert response.status_code == 200
    me = client.get("/api/auth/me", headers=headers(lecturer)).json()
    assert me["must_change_password"] is False


def test_lecturer_creates_course_and_regenerates_enrolment_code():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer)
    assert len(classroom["join_code"]) == 7
    changed = client.post(f"/api/classes/{classroom['id']}/regenerate-code", headers=headers(lecturer))
    assert changed.status_code == 200
    assert changed.json()["join_code"] != classroom["join_code"]


def test_student_registers_and_enrols_with_lecturer_code():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Course for enrolment")
    student = student_account()
    joined = enrol(student, classroom)
    assert joined["id"] == classroom["id"]
    classes = client.get("/api/classes", headers=headers(student)).json()
    assert any(item["id"] == classroom["id"] for item in classes)


def test_course_outline_extracts_objectives_readings_and_subsections():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Structured Course")
    uploaded = upload_outline(lecturer, classroom)
    item = uploaded["uploaded"][0]
    assert item["document_type"] == "course_outline"
    assert item["sections"] >= 4
    assert item["objectives_found"] == 2
    assert item["readings_found"] == 1
    refreshed = client.get("/api/classes", headers=headers(lecturer)).json()
    course = next(row for row in refreshed if row["id"] == classroom["id"])
    assert "Explain descriptive statistics" in course["learning_outcomes"]
    assert any("Introductory Statistics" in item for item in course["recommended_readings"])


def test_student_opens_course_structure_and_generates_grounded_section_lesson():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Section Lesson Course")
    upload_outline(lecturer, classroom)
    student = student_account()
    enrol(student, classroom)
    structure = client.get(f"/api/classes/{classroom['id']}/course-structure", headers=headers(student))
    assert structure.status_code == 200
    documents = structure.json()["documents"]
    sections = [section for document in documents for section in document["sections"]]
    target = next(section for section in sections if "Types of Data" in section["title"])
    lesson = client.post(
        f"/api/course/sections/{target['id']}/teach",
        headers=headers(student),
        json={"level": "University", "detail": "detailed", "include_worked_examples": True, "include_self_check": True},
    )
    assert lesson.status_code == 200, lesson.text
    data = lesson.json()
    assert "Types of Data" in data["section_title"]
    assert data["sources"]
    assert data["visual"]["kind"] == "slides"
    assert data["visual"]["slides"]
    first_slide = data["visual"]["slides"][0]
    assert "explanation" in first_slide
    assert "speaker_note" in first_slide


def test_documents_are_isolated_by_course():
    lecturer, _ = lecturer_account()
    first = create_course(lecturer, name="Class Alpha")
    second = create_course(lecturer, name="Class Beta")
    upload_outline(lecturer, first)
    first_docs = client.get(f"/api/classes/{first['id']}/course-structure", headers=headers(lecturer)).json()["documents"]
    second_docs = client.get(f"/api/classes/{second['id']}/course-structure", headers=headers(lecturer)).json()["documents"]
    assert first_docs
    assert second_docs == []


def test_lecturer_can_upload_teaching_notes_and_recommended_reading():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Multiple Documents")
    for document_type, filename, text in (
        ("teaching_notes", "notes.txt", b"# Unit 1\nDetailed lecturer teaching notes for this unit."),
        ("recommended_reading", "reading.txt", b"# Chapter 2\nApproved recommended reading extract."),
    ):
        response = client.post(
            "/api/materials/upload",
            headers=headers(lecturer),
            data={"class_id": classroom["id"], "document_type": document_type},
            files=[("files", (filename, text, "text/plain"))],
        )
        assert response.status_code == 200, response.text
    structure = client.get(f"/api/classes/{classroom['id']}/course-structure", headers=headers(lecturer)).json()
    types = {item["document_type"] for item in structure["documents"]}
    assert {"teaching_notes", "recommended_reading"}.issubset(types)


def test_required_practice_whiteboard_rejects_typed_only_answer():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Handwritten Practice", required=True)
    student = student_account()
    enrol(student, classroom)
    start = client.post(
        "/api/practice/start",
        headers=headers(student),
        data={"topic": "Mean", "level": "University", "course": "STA 101", "class_id": classroom["id"], "question_count": "2"},
    )
    assert start.status_code == 200
    response = client.post(
        "/api/practice/check",
        headers=headers(student),
        data={"practice_id": start.json()["practice_id"], "answer": "Typed answer only"},
    )
    assert response.status_code == 422
    assert "handwritten" in response.json()["detail"].lower()


def test_required_practice_whiteboard_accepts_image_response():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Handwritten Practice Image", required=True)
    student = student_account()
    enrol(student, classroom)
    start = client.post(
        "/api/practice/start",
        headers=headers(student),
        data={"topic": "Fractions", "level": "Junior High School", "course": "Mathematics", "class_id": classroom["id"], "question_count": "2"},
    ).json()
    response = client.post(
        "/api/practice/check",
        headers=headers(student),
        data={"practice_id": start["practice_id"], "answer": ""},
        files={"board_image": ("practice.png", ONE_PIXEL_PNG, "image/png")},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["completed"] is False
    assert data["question_score"] > 0
    assert data["response_received"] is True


def test_inconsistent_zero_score_cannot_complete_practice():
    evaluation = _normalise_practice_evaluation(
        PracticeEvaluation(correct=True, score=0, feedback="No markable response received")
    )
    assert evaluation.correct is False
    assert evaluation.score == 0


def test_partial_score_is_preserved_for_typed_practice_attempt():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Partial Credit Practice", response_mode="typed")
    student = student_account()
    enrol(student, classroom)
    start = client.post(
        "/api/practice/start",
        headers=headers(student),
        data={"topic": "Mean", "level": "University", "course": "STA 101", "class_id": classroom["id"], "question_count": "2"},
    ).json()
    response = client.post(
        "/api/practice/check",
        headers=headers(student),
        data={"practice_id": start["practice_id"], "answer": "I have started by identifying the values."},
    )
    assert response.status_code == 200, response.text
    data = response.json()
    assert data["completed"] is False
    assert data["question_score"] == 35
    assert data["total_score"] > 0


def test_index_exposes_pause_control_and_capture_status():
    html = client.get("/").text
    assert 'id="pauseTeaching"' in html
    assert 'id="practiceCaptureStatus"' in html
    assert '/static/v2_1.js?v=5.4.0' in html
    assert '/static/practice_board.js?v=5.4.0' in html



def test_demo_chat_returns_visual_plan():
    response = client.post(
        "/api/chat",
        data={
            "message": "Explain photosynthesis.", "session_id": "test-session", "level": "Junior High School",
            "tutor_mode": "guided", "course": "Science", "visual_requested": "true", "visual_preference": "steps",
        },
    )
    assert response.status_code == 200
    data = response.json()
    assert data["demo"] is True
    assert data["visual"]["kind"] == "steps"


def test_demo_work_check_returns_step_level_feedback():
    response = client.post(
        "/api/work/check",
        data={"problem_context": "Solve 2x + 4 = 10", "board_context": '{"ink_strokes":3}', "level": "Junior High School", "course": "Mathematics"},
        files={"board_image": ("working.png", ONE_PIXEL_PNG, "image/png")},
    )
    assert response.status_code == 200
    data = response.json()
    assert 0 <= data["score"] <= 100
    assert "step_results" in data
    assert "first_error_step" in data


def test_visual_slide_schema_keeps_detailed_teaching_fields():
    plan = VisualPlan(
        kind="slides",
        slides=[{
            "title": "Mean", "bullets": ["Definition"], "explanation": "A detailed explanation.",
            "worked_example": "Add the values and divide by their number.", "key_terms": ["sum", "count"],
            "check_question": "What is the mean of 2 and 4?", "speaker_note": "Explain each stage slowly.",
        }],
    )
    slide = plan.slides[0]
    assert slide.explanation
    assert slide.worked_example
    assert slide.check_question


def test_visual_normalisation_clamps_image_boxes():
    plan = VisualPlan(kind="image_annotation", title="Check", annotations=[{"label":"Area","x":990,"y":995,"width":200,"height":100}])
    box = _normalise_visual_plan(plan, has_image=True).annotations[0]
    assert box.x + box.width <= 1000
    assert box.y + box.height <= 1000


def test_image_annotation_becomes_none_without_image():
    plan = VisualPlan(kind="image_annotation", annotations=[{"label":"Area","x":10,"y":10,"width":100,"height":100}])
    assert _normalise_visual_plan(plan, has_image=False).kind == "none"


def test_material_upload_without_course_still_requires_admin_key():
    response = client.post(
        "/api/materials/upload",
        data={"admin_key": "wrong"},
        files={"files": ("notes.txt", b"A simple course note.", "text/plain")},
    )
    assert response.status_code == 401



def test_administrator_documents_are_private_and_deletable():
    admin = admin_account()
    uploaded = client.post(
        "/api/materials/upload",
        headers=headers(admin),
        data={"document_type": "teaching_notes"},
        files={"files": ("admin-private-note.txt", b"Private administrator reference material.", "text/plain")},
    )
    assert uploaded.status_code == 200, uploaded.text
    admin_materials = client.get("/api/admin/materials", headers=headers(admin))
    assert admin_materials.status_code == 200
    item = next(row for row in admin_materials.json()["materials"] if row["source"] == "admin-private-note.txt")
    assert item["source_id"].startswith("global::")

    lecturer, _ = lecturer_account("Lecturer Without Admin Documents")
    classroom = create_course(lecturer, name="Private Repository Isolation")
    lecturer_materials = client.get(f"/api/materials?class_id={classroom['id']}", headers=headers(lecturer))
    assert lecturer_materials.status_code == 200
    assert all(row["source"] != "admin-private-note.txt" for row in lecturer_materials.json()["materials"])

    deleted = client.delete(
        "/api/admin/materials",
        headers=headers(admin),
        params={"source_id": item["source_id"]},
    )
    assert deleted.status_code == 200, deleted.text
    remaining = client.get("/api/admin/materials", headers=headers(admin)).json()["materials"]
    assert all(row["source"] != "admin-private-note.txt" for row in remaining)


def test_lecturer_documents_do_not_cross_lecturer_accounts_and_can_be_deleted():
    lecturer_a, _ = lecturer_account("Lecturer Alpha")
    lecturer_b, _ = lecturer_account("Lecturer Beta")
    class_a = create_course(lecturer_a, name="Alpha Course")
    class_b = create_course(lecturer_b, name="Beta Course")
    upload_outline(lecturer_a, class_a)

    own_structure = client.get(f"/api/classes/{class_a['id']}/course-structure", headers=headers(lecturer_a))
    assert own_structure.status_code == 200
    documents = own_structure.json()["documents"]
    assert documents

    other_structure = client.get(f"/api/classes/{class_b['id']}/course-structure", headers=headers(lecturer_b))
    assert other_structure.status_code == 200
    assert other_structure.json()["documents"] == []

    forbidden = client.get(f"/api/classes/{class_a['id']}/course-structure", headers=headers(lecturer_b))
    assert forbidden.status_code == 403

    document_id = documents[0]["id"]
    deleted = client.delete(f"/api/classes/{class_a['id']}/documents/{document_id}", headers=headers(lecturer_a))
    assert deleted.status_code == 200, deleted.text
    after = client.get(f"/api/classes/{class_a['id']}/course-structure", headers=headers(lecturer_a)).json()["documents"]
    assert after == []
    indexed = client.get(f"/api/materials?class_id={class_a['id']}", headers=headers(lecturer_a)).json()["materials"]
    assert indexed == []


def test_student_sees_every_enrolled_course_from_multiple_lecturers():
    lecturer_a, _ = lecturer_account("Lecturer One")
    lecturer_b, _ = lecturer_account("Lecturer Two")
    class_a = create_course(lecturer_a, name="Enrolled Course One")
    class_b = create_course(lecturer_b, name="Enrolled Course Two")
    student = student_account("Multi-course Student")
    enrol(student, class_a)
    enrol(student, class_b)
    classes = client.get("/api/classes", headers=headers(student))
    assert classes.status_code == 200
    ids = {row["id"] for row in classes.json()}
    assert {class_a["id"], class_b["id"]}.issubset(ids)
    dashboard = client.get("/api/dashboard", headers=headers(student)).json()
    dashboard_ids = {row["id"] for row in dashboard["classes"]}
    assert {class_a["id"], class_b["id"]}.issubset(dashboard_ids)
    assert dashboard["summary"]["classes"] >= 2

def test_legacy_classless_admin_source_is_hidden_and_deletable():
    admin = admin_account()
    source = f"legacy-admin-{uuid.uuid4().hex}.txt"
    chunks = make_chunks(
        "Legacy private administrator content that must never enter lecturer courses.",
        source,
        class_id="",
        material_type="course",
        display_source=source,
        repository_scope="admin_private",
    )
    knowledge.replace_source(source, chunks)

    lecturer, _ = lecturer_account("Legacy Isolation Lecturer")
    classroom = create_course(lecturer, name="Legacy Isolation Course")
    lecturer_rows = client.get(
        f"/api/materials?class_id={classroom['id']}", headers=headers(lecturer)
    ).json()["materials"]
    assert all(row["source_id"] != source for row in lecturer_rows)

    admin_rows = client.get("/api/admin/materials", headers=headers(admin)).json()["materials"]
    assert any(row["source_id"] == source for row in admin_rows)
    deleted = client.delete("/api/admin/materials", headers=headers(admin), params={"source_id": source})
    assert deleted.status_code == 200, deleted.text
    assert deleted.json()["deleted_chunks"] >= 1
    assert knowledge.source_metadata(source) is None


def test_course_delete_removes_legacy_index_aliases():
    lecturer, _ = lecturer_account("Legacy Course Delete Lecturer")
    classroom = create_course(lecturer, name="Legacy Delete Course")
    uploaded = upload_outline(lecturer, classroom)
    document = uploaded["documents"][0]
    legacy_source = document["filename"]
    legacy_chunks = make_chunks(
        "Stale legacy course extract.",
        legacy_source,
        class_id=classroom["id"],
        material_type="course",
        display_source=document["filename"],
        repository_scope="course",
    )
    knowledge.replace_source(legacy_source, legacy_chunks)
    response = client.delete(
        f"/api/classes/{classroom['id']}/documents/{document['id']}",
        headers=headers(lecturer),
    )
    assert response.status_code == 200, response.text
    rows = client.get(f"/api/materials?class_id={classroom['id']}", headers=headers(lecturer)).json()["materials"]
    assert rows == []


def test_portal_and_api_responses_disable_stale_caching():
    root_response = client.get("/")
    api_response = client.get("/api/config")
    assert "no-store" in root_response.headers.get("cache-control", "")
    assert "no-store" in api_response.headers.get("cache-control", "")


def test_extract_response_text_from_message_items():
    class Part:
        type = "output_text"
        text = "Visible answer"
    class Item:
        type = "message"
        content = [Part()]
    class FakeResponse:
        output_text = ""
        output = [Item()]
    assert _extract_response_text(FakeResponse()) == "Visible answer"


def test_audio_mime_variants_are_accepted():
    assert _base_media_type("audio/webm;codecs=opus") == "audio/webm"
    assert _audio_upload_extension("question.webm", "audio/webm;codecs=opus") == ".webm"
    assert _audio_upload_extension("question.bin", "audio/mp4;codecs=mp4a.40.2") == ".m4a"
    assert _audio_upload_extension("question.webm", "application/octet-stream") == ".webm"


def test_service_worker_and_manifest_are_v5():
    manifest = client.get("/static/manifest.webmanifest")
    worker = client.get("/static/service-worker.js")
    assert manifest.status_code == 200
    assert worker.status_code == 200
    assert "Anovlad Institutional AI Tutor" in manifest.text
    assert "anovlad-ai-tutor-v5-4-0-shell" in worker.text


def test_cost_aware_router_prefers_flash_for_normal_and_pro_for_advanced():
    from app.main import ai_router
    normal_model, _ = ai_router.choose_deepseek_model("Explain photosynthesis simply.")
    advanced_model, _ = ai_router.choose_deepseek_model("Derive an advanced stochastic differential equation and prove the theorem using eigenvalues.")
    assert normal_model == "deepseek-v4-flash"
    assert advanced_model == "deepseek-v4-pro"


def weekly_outline_docx_bytes():
    output = BytesIO()
    document = Document()
    document.add_heading("STA 102 Weekly Course Outline", 0)
    document.add_heading("Course Objectives", 1)
    document.add_paragraph("1. Explain the role of statistics in decision making")
    document.add_paragraph("2. Classify data correctly")
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Week"
    table.rows[0].cells[1].text = "Topic"
    table.rows[0].cells[2].text = "Activities"
    rows = [
        ("1", "Introduction to Statistics", "Meaning of statistics; Uses of statistics; Limitations"),
        ("2", "Types of Data", "Qualitative data; Quantitative data; Levels of measurement"),
    ]
    for week, topic, activities in rows:
        cells = table.add_row().cells
        cells[0].text = week
        cells[1].text = topic
        cells[2].text = activities
    document.save(output)
    return output.getvalue()


def test_v51_student_workspace_controls_are_present():
    html = client.get("/").text
    for identifier in (
        'id="practiceResponseChooser"', 'id="practiceRecordResponse"',
        'id="practiceAudioPreview"', 'id="practiceAddSpace"',
        'id="practiceFullscreen"', 'id="fullscreenBoard"',
    ):
        assert identifier in html
    assert '/static/v2_1.js?v=5.4.0' in html
    css = client.get('/static/styles.css').text
    assert 'body.student-interface' in css
    assert '.practice-whiteboard-wrap:fullscreen' in css
    assert 'overflow-y: auto' in css


def test_lecturer_can_require_each_practice_response_mode():
    lecturer, _ = lecturer_account()
    student = student_account()
    for mode in ("typed", "voice", "whiteboard"):
        classroom = create_course(lecturer, name=f"{mode.title()} Practice Course", response_mode=mode)
        enrol(student, classroom)
        response = client.post(
            "/api/practice/start",
            headers=headers(student),
            data={
                "topic": "Measures of central tendency",
                "course": classroom["name"],
                "level": "University",
                "question_count": 2,
                "class_id": classroom["id"],
            },
        )
        assert response.status_code == 200, response.text
        question_data = response.json()
        assert question_data["response_mode"] == mode
        assert question_data["allowed_response_modes"] == [mode]
        empty_check = client.post(
            "/api/practice/check",
            headers=headers(student),
            data={"practice_id": question_data["practice_id"], "answer": ""},
        )
        assert empty_check.status_code == 422


def test_student_choice_practice_allows_typing_voice_and_whiteboard():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Choice Practice Course", response_mode="student_choice")
    student = student_account()
    enrol(student, classroom)
    response = client.post(
        "/api/practice/start",
        headers=headers(student),
        data={"topic": "Sampling", "class_id": classroom["id"], "question_count": 2},
    )
    assert response.status_code == 200
    assert response.json()["response_mode"] == "student_choice"
    assert response.json()["allowed_response_modes"] == ["typed", "voice", "whiteboard"]


def test_weekly_course_outline_table_displays_weeks_and_selectable_subunits():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Weekly Outline Course")
    upload = client.post(
        "/api/materials/upload",
        headers=headers(lecturer),
        data={"class_id": classroom["id"], "document_type": "course_outline"},
        files=[(
            "files",
            ("weekly-outline.docx", weekly_outline_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        )],
    )
    assert upload.status_code == 200, upload.text
    student = student_account()
    enrol(student, classroom)
    structure = client.get(f"/api/classes/{classroom['id']}/course-structure", headers=headers(student))
    assert structure.status_code == 200
    plan = structure.json()["weekly_plan"]
    assert len(plan) >= 2
    assert plan[0]["title"].startswith("Week 1")
    assert plan[0]["subunits"]
    first_subunit = plan[0]["subunits"][0]
    assert first_subunit["id"]
    lesson = client.post(
        f"/api/course/sections/{first_subunit['id']}/teach",
        headers=headers(student),
        json={"level": "University", "detail": "detailed"},
    )
    assert lesson.status_code == 200, lesson.text
    assert lesson.json()["visual"]["kind"] == "slides"


def test_course_without_readings_generates_weekly_lessons_from_outcomes():
    lecturer, _ = lecturer_account()
    classroom = create_course(
        lecturer,
        name="Outcome Generated Course",
        weekly_topics=["Foundations of research", "Developing research questions"],
        outcomes=["Explain the research process", "Formulate answerable research questions"],
    )
    student = student_account()
    enrol(student, classroom)
    structure = client.get(f"/api/classes/{classroom['id']}/course-structure", headers=headers(student)).json()
    assert len(structure["weekly_plan"]) == 2
    assert all(item["generated"] for item in structure["weekly_plan"])
    topic = structure["weekly_plan"][0]
    lesson = client.post(
        f"/api/course/sections/{topic['id']}/teach",
        headers=headers(student),
        json={"level": "University", "detail": "extended"},
    )
    assert lesson.status_code == 200, lesson.text
    data = lesson.json()
    assert data["generated_from_outcomes"] is True
    assert data["sources"] == ["Lecturer course objectives, expected outcomes and weekly plan"]
    assert data["visual"]["kind"] == "slides"
    assert data["answer"]


def period_based_operations_outline_docx_bytes():
    output = BytesIO()
    document = Document()
    document.add_paragraph("2025/2026 FIRST SEMESTER")
    info = document.add_table(rows=2, cols=2)
    info.rows[0].cells[0].text = "Name of Lecturer(s)"
    info.rows[0].cells[1].text = "Lecturer Name"
    info.rows[1].cells[0].text = "Course Code/ Title"
    info.rows[1].cells[1].text = "SBU301: Operations Management"
    document.add_paragraph("2.0 Course Description:")
    document.add_paragraph("The course develops understanding of operations as a core business function.")
    document.add_paragraph("3.0 Course Objectives:")
    document.add_paragraph("This course enables students to:")
    document.add_paragraph("acquire knowledge in operations strategy and organisational competitiveness.", style="List Paragraph")
    document.add_paragraph("use quantitative approaches to make operations decisions.", style="List Paragraph")
    document.add_paragraph("3.0 Course Outcomes:")
    document.add_paragraph("Use equations and statistics in operations analysis", style="List Paragraph")
    document.add_heading("4.0 Course Outline", level=3)
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Period"
    table.rows[0].cells[1].text = "Topics"
    table.rows[0].cells[2].text = "Student’s Preparation"
    rows = [
        (
            "One",
            ["Introduction to operations management", "meaning of operations management", "why study OM", "what operations managers do"],
            ["Read chapter 1", "Prepare an analysis of how operations, marketing and finance are related."],
        ),
        (
            "Two",
            ["Operations strategy in a global environment", "developing missions and strategies", "achieving competitive advantage through operations"],
            ["Read pages 68-96", "Analyse the mission and strategy of an organisation."],
        ),
    ]
    for period, topics, preparation in rows:
        cells = table.add_row().cells
        cells[0].text = period
        cells[1].text = topics[0]
        for item in topics[1:]:
            cells[1].add_paragraph(item, style="List Paragraph")
        cells[2].text = preparation[0]
        for item in preparation[1:]:
            cells[2].add_paragraph(item)
    document.save(output)
    return output.getvalue()


def test_period_topics_preparation_outline_exposes_weeks_subtopics_and_activities():
    lecturer, _ = lecturer_account()
    classroom = create_course(lecturer, name="Operations Management")
    upload = client.post(
        "/api/materials/upload",
        headers=headers(lecturer),
        data={"class_id": classroom["id"], "document_type": "course_outline"},
        files=[(
            "files",
            (
                "Course Outline Operations Management.docx",
                period_based_operations_outline_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        )],
    )
    assert upload.status_code == 200, upload.text
    uploaded = upload.json()["uploaded"][0]
    assert uploaded["weekly_topics_found"] == 2
    assert uploaded["objectives_found"] >= 3

    student = student_account()
    enrol(student, classroom)
    structure = client.get(
        f"/api/classes/{classroom['id']}/course-structure",
        headers=headers(student),
    )
    assert structure.status_code == 200, structure.text
    data = structure.json()
    assert data["documents"][0]["title"] == "SBU301: Operations Management"
    plan = data["weekly_plan"]
    assert [item["title"] for item in plan] == [
        "Week 1: Introduction to operations management",
        "Week 2: Operations strategy in a global environment",
    ]
    assert [item["title"] for item in plan[0]["subunits"]] == [
        "meaning of operations management",
        "why study OM",
        "what operations managers do",
    ]
    assert plan[0]["preparation"] == [
        "Read chapter 1",
        "Prepare an analysis of how operations, marketing and finance are related.",
    ]


def test_guided_lecture_speech_request_supports_style_and_natural_speed():
    request = SpeechRequest(
        text="Explain the concept carefully.",
        voice="nova",
        style="guided_lecture",
        speed=0.94,
    )
    assert request.style == "guided_lecture"
    assert request.speed == 0.94


def test_health_exposes_guided_lecture_capabilities():
    data = client.get("/health").json()
    assert data["guided_lecture_notes_enabled"] is True
    assert data["synchronised_slide_popups_enabled"] is True
    assert data["natural_lecture_pacing_enabled"] is True


def test_frontend_uses_continuous_guided_lecture_audio_and_progressive_popups():
    root = os.path.join(os.path.dirname(__file__), "..", "app", "static")
    v2 = open(os.path.join(root, "v2_1.js"), encoding="utf-8").read()
    app_js = open(os.path.join(root, "app.js"), encoding="utf-8").read()
    css = open(os.path.join(root, "styles.css"), encoding="utf-8").read()
    assert "style: 'guided_lecture'" in v2
    assert "speed: 0.94" in v2
    assert "playLectureChunk" in v2
    assert "data-lecture-cue" in app_js
    assert "lecture-revealed" in css
    assert "lecture-notes-panel" in css


def test_v54_health_and_frontend_expose_complete_learning_cycle():
    health = client.get('/health').json()
    assert health['version'] == '5.4.0'
    for key in (
        'diagnostic_mastery_engine_enabled', 'personalised_learning_paths_enabled',
        'lecturer_assessment_manager_enabled', 'intelligent_remediation_enabled',
        'spaced_revision_enabled', 'student_notes_bookmarks_enabled',
        'academic_integrity_controls_enabled', 'accessibility_controls_enabled',
    ):
        assert health[key] is True
    html = client.get('/').text
    assert '/static/portal.js?v=5.4.0' in html
    assert 'data-lesson-followup=' in html
    assert 'id="repeatLastExplanation"' in html
    portal = client.get('/static/portal.js').text
    for marker in ('Your personalised learning home', 'Assessment and question bank', 'Due for review', 'My notes and bookmarks', 'Remedial lesson'):
        assert marker in portal


def test_entry_diagnostic_personalised_path_mastery_and_revision_cycle():
    lecturer, _ = lecturer_account('Diagnostic Lecturer')
    classroom = create_course(
        lecturer,
        name='Diagnostic Learning Course',
        weekly_topics=['Foundations', 'Applications'],
        outcomes=['Explain the foundations', 'Apply the core method'],
    )
    student = student_account('Diagnostic Student')
    enrol(student, classroom)

    dashboard = client.get('/api/dashboard', headers=headers(student))
    assert dashboard.status_code == 200, dashboard.text
    data = dashboard.json()
    diagnostics = [item for item in data['assessments'] if item['assessment_type'] == 'diagnostic']
    assert diagnostics
    assert data['next_recommended_action']['type'] == 'diagnostic'
    assert data['learning_paths'][0]['diagnostic_required'] is True

    started = client.post(f"/api/assessments/{diagnostics[0]['id']}/start", headers=headers(student))
    assert started.status_code == 200, started.text
    attempt = started.json()
    questions = attempt['assessment']['questions']
    assert questions
    assert all('expected_answer' not in question for question in questions)
    responses = [
        {'question_id': question['id'], 'answer': 'I understand the main idea, can define the concept accurately, and can give a relevant practical example that shows how it is applied.', 'mode': 'typed'}
        for question in questions
    ]
    submitted = client.post(
        f"/api/assessment-attempts/{attempt['attempt_id']}/submit",
        headers=headers(student),
        json={'responses': responses, 'hints_used': 0},
    )
    assert submitted.status_code == 200, submitted.text
    result = submitted.json()
    assert result['score'] >= 0
    assert result['next_action']

    refreshed = client.get('/api/dashboard', headers=headers(student)).json()
    assert refreshed['learning_paths'][0]['diagnostic_required'] is False
    assert refreshed['mastery_records']
    assert refreshed['reviews_due']
    assert refreshed['next_recommended_action']['type'] in {'revision', 'lesson', 'course_complete'}


def test_lecturer_can_generate_edit_publish_and_student_complete_assessment():
    lecturer, _ = lecturer_account('Assessment Lecturer')
    classroom = create_course(
        lecturer,
        name='Assessment Design Course',
        weekly_topics=['Demand forecasting'],
        outcomes=['Apply a forecasting method'],
    )
    draft = client.post(
        f"/api/classes/{classroom['id']}/assessments/draft",
        headers=headers(lecturer),
        json={'assessment_type': 'quiz', 'topic': 'Demand forecasting', 'learning_outcome': 'Apply a forecasting method', 'question_count': 3, 'difficulty': 'mixed'},
    )
    assert draft.status_code == 200, draft.text
    item = draft.json()
    assert item['status'] == 'draft'
    assert len(item['questions']) == 3
    item['title'] = 'Lecturer-reviewed forecasting quiz'
    item['instructions'] = 'Show the reasoning behind each answer.'
    item['settings'] = {
        'attempts_allowed': 2, 'hints_allowed': True, 'reveal_answers': True,
        'pass_mark': 65, 'contributes_to_mastery': True,
        'integrity_mode': 'graded', 'deadline_enforced': False,
    }
    item['status'] = 'published'
    updated = client.patch(f"/api/assessments/{item['id']}", headers=headers(lecturer), json=item)
    assert updated.status_code == 200, updated.text
    assert updated.json()['title'] == 'Lecturer-reviewed forecasting quiz'
    assert updated.json()['settings']['attempts_allowed'] == 2

    student = student_account('Assessment Student')
    enrol(student, classroom)
    listed = client.get(f"/api/classes/{classroom['id']}/assessments", headers=headers(student))
    assert listed.status_code == 200
    quiz = next(row for row in listed.json() if row['id'] == item['id'])
    assert all('marking_guide' not in question for question in quiz['questions'])
    attempt = client.post(f"/api/assessments/{item['id']}/start", headers=headers(student)).json()
    responses = [{'question_id': question['id'], 'answer': 'This answer explains the method, identifies the relevant variables, applies the steps and gives a practical example.', 'mode': 'typed'} for question in attempt['assessment']['questions']]
    completed = client.post(f"/api/assessment-attempts/{attempt['attempt_id']}/submit", headers=headers(student), json={'responses': responses, 'hints_used': 0})
    assert completed.status_code == 200, completed.text
    assert 'mastery' in completed.json()


def test_student_notes_are_private_and_revision_sheet_is_available():
    lecturer, _ = lecturer_account('Notes Lecturer')
    classroom = create_course(lecturer, name='Notes Course')
    student_one = student_account('Notes Student One')
    student_two = student_account('Notes Student Two')
    enrol(student_one, classroom)
    enrol(student_two, classroom)
    created = client.post('/api/student/notes', headers=headers(student_one), json={
        'class_id': classroom['id'], 'section_id': 'week-1', 'note_type': 'bookmark',
        'title': 'Important definition', 'content': 'My private summary of the definition.', 'metadata': {},
    })
    assert created.status_code == 200, created.text
    note = created.json()
    own = client.get('/api/student/notes', headers=headers(student_one)).json()
    other = client.get('/api/student/notes', headers=headers(student_two)).json()
    assert any(item['id'] == note['id'] for item in own)
    assert all(item['id'] != note['id'] for item in other)
    sheet = client.get(f"/api/student/revision-sheet?class_id={classroom['id']}", headers=headers(student_one))
    assert sheet.status_code == 200
    assert 'Important definition' in sheet.text
    deleted = client.delete(f"/api/student/notes/{note['id']}", headers=headers(student_one))
    assert deleted.status_code == 200


def test_course_learning_and_integrity_policy_is_exposed_to_enrolled_students():
    lecturer, _ = lecturer_account('Integrity Lecturer')
    classroom = create_course(lecturer, name='Integrity Course')
    payload = {
        'name': classroom['name'], 'subject': classroom['subject'], 'knowledge_mode': 'course_only',
        'learning_outcomes': ['Explain ethical use of AI'], 'weekly_topics': ['Academic integrity'],
        'recommended_readings': [], 'tutor_instructions': 'Use hints before direct answers.',
        'practice_whiteboard_required': False, 'practice_response_mode': 'student_choice',
        'diagnostics_required': True, 'spaced_revision_enabled': True, 'mastery_pass_mark': 75,
        'direct_answers_allowed': False, 'hints_allowed': True,
        'assignment_help_mode': 'teach_only', 'integrity_mode': 'hint_only',
    }
    updated = client.patch(f"/api/classes/{classroom['id']}/profile", headers=headers(lecturer), json=payload)
    assert updated.status_code == 200, updated.text
    student = student_account('Integrity Student')
    joined = enrol(student, updated.json())
    assert joined['direct_answers_allowed'] is False
    assert joined['assignment_help_mode'] == 'teach_only'
    assert joined['integrity_mode'] == 'hint_only'
    assert joined['mastery_pass_mark'] == 75


def test_teacher_dashboard_flags_enrolled_student_with_no_activity():
    lecturer, _ = lecturer_account('Intervention Lecturer')
    classroom = create_course(lecturer, name='Intervention Course')
    student = student_account('Inactive Enrolled Student')
    enrol(student, classroom)
    dashboard = client.get('/api/dashboard', headers=headers(lecturer))
    assert dashboard.status_code == 200, dashboard.text
    intervention = next(item for item in dashboard.json()['interventions'] if item.get('id') == student['user']['id'])
    assert 'No learning activity recorded' in intervention['reasons']
    assert intervention['recommended_action']


def test_assessment_response_document_extraction_and_private_student_access():
    student = student_account('Response Extraction Student')
    extracted = client.post(
        '/api/assessment/response/extract',
        headers=headers(student),
        data={'response_mode': 'upload'},
        files={'file': ('answer.txt', b'My structured written response with supporting explanation.', 'text/plain')},
    )
    assert extracted.status_code == 200, extracted.text
    assert 'structured written response' in extracted.json()['text']
    lecturer, _ = lecturer_account('No Extraction Lecturer')
    forbidden = client.post(
        '/api/assessment/response/extract',
        headers=headers(lecturer),
        data={'response_mode': 'upload'},
        files={'file': ('answer.txt', b'Lecturer file', 'text/plain')},
    )
    assert forbidden.status_code == 403


def test_assessment_deadline_and_hidden_hints_are_enforced():
    lecturer, _ = lecturer_account('Deadline Lecturer')
    classroom = create_course(lecturer, name='Deadline Course')
    created = client.post(
        f"/api/classes/{classroom['id']}/assessments",
        headers=headers(lecturer),
        json={
            'title': 'Closed timed quiz', 'assessment_type': 'quiz', 'topic': 'Topic',
            'learning_outcome': 'Outcome', 'instructions': 'Answer independently.',
            'questions': [{
                'id': 'q1', 'question_type': 'short_answer', 'prompt': 'Explain the concept.',
                'expected_answer': 'Expected answer', 'marking_guide': 'Marking guide',
                'hint': 'Private hint', 'explanation': 'Private explanation', 'difficulty': 'standard',
                'points': 1, 'response_mode': 'typed', 'options': [],
            }],
            'settings': {
                'attempts_allowed': 1, 'hints_allowed': False, 'reveal_answers': True,
                'pass_mark': 70, 'contributes_to_mastery': True,
                'integrity_mode': 'exam', 'deadline_enforced': True,
            },
            'status': 'published', 'due_at': '2020-01-01T00:00:00+00:00',
        },
    )
    assert created.status_code == 200, created.text
    student = student_account('Deadline Student')
    enrol(student, classroom)
    listed = client.get(f"/api/classes/{classroom['id']}/assessments", headers=headers(student)).json()
    question = next(item for item in listed if item['id'] == created.json()['id'])['questions'][0]
    assert 'hint' not in question
    assert 'expected_answer' not in question
    assert 'explanation' not in question
    start = client.post(f"/api/assessments/{created.json()['id']}/start", headers=headers(student))
    assert start.status_code == 409
    assert 'deadline' in start.json()['detail'].lower()


def test_diagnostic_records_question_level_mastery_for_outcomes_and_topics():
    lecturer, _ = lecturer_account('Mapped Diagnostic Lecturer')
    outcomes = ['Explain the core concept', 'Apply the core concept']
    topics = ['Foundations of the course', 'Practical applications']
    classroom = create_course(
        lecturer,
        name='Mapped Diagnostic Course',
        weekly_topics=topics,
        outcomes=outcomes,
    )
    student = student_account('Mapped Diagnostic Student')
    enrol(student, classroom)
    dashboard = client.get('/api/dashboard', headers=headers(student)).json()
    diagnostic = next(item for item in dashboard['assessments'] if item['assessment_type'] == 'diagnostic')
    attempt = client.post(f"/api/assessments/{diagnostic['id']}/start", headers=headers(student)).json()
    question_map = {
        question['id']: (question.get('learning_outcome', ''), question.get('topic', ''))
        for question in attempt['assessment']['questions']
    }
    assert any(outcome for outcome, _ in question_map.values())
    assert any(topic for _, topic in question_map.values())
    responses = [
        {
            'question_id': question['id'],
            'answer': f"I can explain {question.get('learning_outcome') or question.get('topic')} accurately and give a practical example.",
            'mode': 'typed',
        }
        for question in attempt['assessment']['questions']
    ]
    submitted = client.post(
        f"/api/assessment-attempts/{attempt['attempt_id']}/submit",
        headers=headers(student),
        json={'responses': responses, 'hints_used': 0},
    )
    assert submitted.status_code == 200, submitted.text
    mastery = submitted.json()['mastery']
    assert mastery['assessment_type'] == 'diagnostic'
    assert mastery['records']
    recorded_keys = {
        item.get('learning_outcome') or item.get('topic')
        for item in mastery['records']
    }
    assert set(outcomes + topics).issubset(recorded_keys)
    path = client.get(f"/api/learning-path/{classroom['id']}", headers=headers(student)).json()
    assert all(item['status'] != 'not_started' for item in path['items'])


def test_mastery_certificate_is_blocked_before_full_mastery():
    lecturer, _ = lecturer_account('Certificate Lecturer')
    classroom = create_course(
        lecturer,
        name='Certificate Course',
        weekly_topics=['Topic one'],
        outcomes=['Master the topic'],
    )
    student = student_account('Certificate Student')
    enrol(student, classroom)
    certificate = client.get(f"/api/student/certificate/{classroom['id']}", headers=headers(student))
    assert certificate.status_code == 409
    assert 'master' in certificate.json()['detail'].lower()
